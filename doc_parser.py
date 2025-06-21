from pathlib import Path
from typing import List
import docx

from utils import normalize_whitespace


def extract_text(doc_path: str | Path) -> str:
    """Extract full text (including tables) from a .docx file."""
    path = Path(doc_path)
    if not path.exists():
        raise FileNotFoundError(path)

    doc = docx.Document(str(path))
    texts: List[str] = []

    # Paragraphs
    for p in doc.paragraphs:
        texts.append(p.text)

    # Tables
    for tbl in doc.tables:
        for row in tbl.rows:
            row_text = " ".join(cell.text for cell in row.cells)
            texts.append(row_text)

    return normalize_whitespace("\n".join(texts)) 